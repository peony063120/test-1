import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

SRC = r"E:\Lap_trinh_nang_cao\project-update-report.docx"
OUT = r"E:\Lap_trinh_nang_cao\project-update-report-updated.docx"

doc = Document(SRC)

# ---------- helpers ----------
def add_heading(text, size=14, bold=True, color=(0x1F, 0x4E, 0x79)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*color)
    return p

def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p

def add_code(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
    run.font.size = Pt(9)
    return p

def add_table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return t

# =========================================================
# SECTION A: LAN nội bộ — giải pháp được chọn
# (inserted BEFORE the "Server cache" section)
# =========================================================
server_cache_par = None
for p in doc.paragraphs:
    if p.text.strip() == 'Server cache':
        server_cache_par = p
        break

# Build content paragraphs to insert before "Server cache"
# We'll insert by adding elements before the found paragraph.
def insert_paragraph_before(ref_par, text, bold=False, size=11):
    new_p = ref_par.insert_paragraph_before()
    run = new_p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return new_p

def insert_bullet_before(ref_par, text, size=11):
    new_p = ref_par.insert_paragraph_before()
    new_p.style = doc.styles['List Bullet']
    run = new_p.add_run(text)
    run.font.size = Pt(size)
    return new_p

def insert_code_before(ref_par, text):
    new_p = ref_par.insert_paragraph_before()
    run = new_p.add_run(text)
    run.font.name = 'Consolas'
    run._element.rPr.rFonts.set(qn('w:ascii'), 'Consolas')
    run.font.size = Pt(9)
    return new_p

if server_cache_par is not None:
    ref = server_cache_par

    # Heading first
    p = ref.insert_paragraph_before()
    r = p.add_run("LAN nội bộ — Giải pháp được chọn")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    # Body bullets (insert in original order → stacked correctly)
    for txt in [
        "Kết hợp 2 kỹ thuật:",
    ]:
        insert_paragraph_before(ref, txt, bold=False)

    for txt in [
        "1. Vite bind 0.0.0.0 (host: true) — cho mọi thiết bị trong LAN truy cập frontend (port 3001).",
        "2. Auto-detect backend URL theo hostname — frontend đọc IP từ URL đang truy cập và tự gọi backend cùng IP, khác port (3000).",
        "3. Backend CORS mở cho mọi LAN origin (192.168.*, 10.*, 172.16-31.*, localhost).",
    ]:
        insert_bullet_before(ref, txt)

    insert_paragraph_before(ref, "Code auto-detect trong axios.config.ts:", bold=False)

    # Code block (insert in original order)
    code_lines = [
        "const resolveBaseUrl = () => {",
        "  const explicit = import.meta.env.VITE_API_URL;",
        "  if (explicit && explicit.trim()) return explicit.trim();",
        "  const { hostname, protocol } = window.location;",
        "  if (hostname === 'localhost' || hostname === '127.0.0.1') {",
        "    return '/api/v1';                      // may chu: Vite proxy",
        "  }",
        "  return `${protocol}//${hostname}:3000/api/v1`;  // LAN IP: goi thang backend",
        "};",
    ]
    for line in code_lines:
        insert_code_before(ref, line)

    # Results
    insert_paragraph_before(ref, "Kết quả kiểm chứng:", bold=True)
    for txt in [
        "localhost:3001 → hoạt động (qua Vite proxy)",
        "http://192.168.100.41:3001 → login thành công, redirect /admin",
        "http://172.28.144.1:3001 → login thành công, redirect /admin",
        "Backend qua LAN IP (192.168.100.41:3000) → ADMIN login OK",
        "Lý do chọn: không cần cấu hình khi đổi mạng; không phụ thuộc dịch vụ ngoài; không đổi data structure.",
    ]:
        insert_bullet_before(ref, txt)

# =========================================================
# SECTION B: Server cache — Giải pháp được chọn (append at end)
# =========================================================
doc.add_paragraph()
add_heading("Server cache — Giải pháp được chọn")
add_para("Chiến lược chọn: Cache-Aside (Lazy Loading) + TTL + Invalidation chủ động", bold=True)
add_bullet("Nghiệp vụ đọc nhiều hơn ghi (quét barcode, xem sản phẩm, tồn kho, báo cáo).")
add_bullet("An toàn dữ liệu: DB vẫn là source of truth, cache lỗi không làm mất dữ liệu.")
add_bullet("Code đã theo pattern này sẵn — chỉ thay in-memory Map bằng Redis thật (ioredis), không đổi kiến trúc.")
add_bullet("Kết hợp invalidation chủ động khi write: update/delete sản phẩm → xóa product:{id}; điều chỉnh tồn kho → xóa inventory:{...}; đổi quyền → xóa permissions:*.")

add_para("Nâng cấp RedisService: ioredis + graceful fallback", bold=True)
add_bullet("Kết nối Redis thật qua REDIS_HOST/REDIS_PORT (từ .env).")
add_bullet("Graceful fallback: Redis down → dùng in-memory Map, app vẫn chạy.")
add_bullet("Wildcard delete (del('permissions:*')) dùng SCAN + DEL — sửa lỗi của bản in-memory cũ.")
add_bullet("Tự disconnect khi app shutdown (onModuleDestroy).")

add_para("Dữ liệu cache theo độ ưu tiên", bold=True)
add_table(
    ["Dữ liệu", "Key", "TTL", "Invalidate khi"],
    [
        ["Permissions theo user", "permissions:{userId}", "3600s", "Đổi role/user, logout"],
        ["Product theo ID", "product:{id}", "300s", "Update/delete product"],
        ["Product theo barcode", "barcode:{code}", "600s", "Update product"],
        ["Inventory theo product", "inventory:{productId}:{warehouseId}", "60s", "Adjust inventory"],
    ],
)

add_para("Kết quả kiểm chứng", bold=True)
add_bullet("Kết nối Redis: 'Redis connected at localhost:6379'.")
add_bullet("Dữ liệu lưu thật trong Redis: redis-cli đọc được giá trị.")
add_bullet("Cache round-trip (set → get): hoạt động.")
add_bullet("Wildcard delete (permissions:*): xóa đúng các key khớp.")
add_bullet("Graceful fallback (Redis down): không crash, dùng in-memory.")
add_bullet("Backend test suite: 21/21 pass, 61 tests.")

doc.save(OUT)
print("Saved:", OUT)
